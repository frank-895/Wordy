"""
Document processing service for handling different file types
"""
import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from typing import Optional, List, Dict, Any
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Service for processing different document types and extracting text content
    """
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text content from a PDF file
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text content from a DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def process_document(file_path: Optional[str] = None, content: Optional[str] = None, 
                        file_type: str = 'text') -> str:
        """
        Process a document and extract its text content
        
        Args:
            file_path: Path to the file (for file uploads)
            content: Direct text content (for text input)
            file_type: Type of document ('text', 'pdf', 'docx')
            
        Returns:
            Extracted text content
        """
        if file_type == 'text':
            if content is None:
                raise ValueError("Content is required for text type documents")
            return content.strip()
        
        elif file_type == 'pdf':
            if file_path is None:
                raise ValueError("File path is required for PDF documents")
            if not os.path.exists(file_path):
                raise ValueError(f"PDF file not found: {file_path}")
            return DocumentProcessor.extract_text_from_pdf(file_path)
        
        elif file_type == 'docx':
            if file_path is None:
                raise ValueError("File path is required for DOCX documents")
            if not os.path.exists(file_path):
                raise ValueError(f"DOCX file not found: {file_path}")
            return DocumentProcessor.extract_text_from_docx(file_path)
        
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def validate_file_type(file_path: str) -> str:
        """
        Validate and determine file type based on file extension
        
        Args:
            file_path: Path to the file
            
        Returns:
            File type ('pdf', 'docx', 'text')
        """
        if not file_path:
            return 'text'
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return 'pdf'
        elif file_extension == '.docx':
            return 'docx'
        else:
            raise ValueError(f"Unsupported file extension: {file_extension}") 