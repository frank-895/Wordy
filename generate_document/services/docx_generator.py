from docx import Document
import io

def build_docx(blocks):
    """
    Builds a DOCX file from structured blocks.
    Supports: heading, paragraph, quote, code, list.
    """
    doc = Document()

    for block in blocks:
        block_type = block[0]

        if block_type == 'heading':
            text, level = block[1], block[2]
            doc.add_heading(text, level=level)

        elif block_type == 'paragraph':
            doc.add_paragraph(block[1])

        elif block_type == 'quote':
            doc.add_paragraph(block[1], style='Intense Quote')

        elif block_type == 'code':
            text, language = block[1], block[2]
            doc.add_paragraph(text, style='Intense Quote')  # Consider custom code style here

        elif block_type == 'list':
            items, list_type = block[1], block[2]
            style = 'ListBullet' if list_type == 'bullet' else 'ListNumber'
            for item in items:
                doc.add_paragraph(item, style=style)

        else:
            # Optional: fallback for unknown types
            doc.add_paragraph(f"[Unsupported block type: {block_type}]")

    # Finalize document
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
